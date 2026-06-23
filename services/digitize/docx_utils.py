"""
Utilities for processing DOCX files, providing page count estimation and TOC extraction.
This module provides DOCX-specific functionality parallel to PDF utilities.
"""
import json
import time
import re
from pathlib import Path
from typing import Dict, List, Tuple, Optional
from docx import Document

from common.misc_utils import get_logger
from common.llm_utils import tqdm_wrapper

logger = get_logger("docx_utils")

# Constants for page estimation
CHARS_PER_PAGE = 3000  # Approximate characters per page
WORDS_PER_PAGE = 450   # Approximate words per page

TABLE_CAPTION_PATTERN = re.compile(
    r"^\s*table\s+\d+(?:[.-]\d+)*\s*[:.-]?\s+.+$",
    re.IGNORECASE
)


def _parse_ref_index(ref: str, prefix: str) -> int | None:
    """
    Parse a Docling ref like '#/texts/616' or '#/tables/2' into its integer index.
    """
    try:
        expected_prefix = f"#/{prefix}/"
        if not isinstance(ref, str) or not ref.startswith(expected_prefix):
            logger.debug(f"_parse_ref_index: ref '{ref}' does not match expected prefix '{expected_prefix}'")
            return None
        parsed_idx = int(ref[len(expected_prefix):])
        logger.debug(f"_parse_ref_index: parsed ref '{ref}' with prefix '{prefix}' to index {parsed_idx}")
        return parsed_idx
    except Exception as e:
        logger.debug(f"_parse_ref_index: failed to parse ref '{ref}' with prefix '{prefix}': {e}")
        return None


def _get_body_children_refs(converted_doc) -> list[str]:
    """
    Return top-level body child refs in document order.
    """
    try:
        children = getattr(converted_doc.body, "children", []) or []
        refs = []
        logger.debug(f"_get_body_children_refs: found {len(children)} top-level body children")
        for idx, child in enumerate(children):
            if isinstance(child, dict) and "$ref" in child:
                refs.append(child["$ref"])
            else:
                child_ref = (
                    getattr(child, "ref", None)
                    or getattr(child, "$ref", None)
                    or getattr(child, "cref", None)
                    or getattr(child, "self_ref", None)
                )
                if not child_ref and hasattr(child, "__dict__"):
                    logger.debug(
                        f"_get_body_children_refs: child[{idx}] available attrs={list(vars(child).keys())}, type={type(child)}"
                    )
                if child_ref:
                    refs.append(child_ref)
        logger.debug(f"_get_body_children_refs: extracted {len(refs)} refs")
        return refs
    except Exception as e:
        logger.debug(f"_get_body_children_refs: failed to extract body children refs: {e}", exc_info=True)
        return []


def _get_text_value_by_ref(converted_doc, ref: str) -> str:
    """
    Resolve a '#/texts/<n>' ref to its text content.
    """
    idx = _parse_ref_index(ref, "texts")
    if idx is None:
        logger.debug(f"_get_text_value_by_ref: could not parse text ref '{ref}'")
        return ""

    try:
        text_obj = converted_doc.texts[idx]
        text = getattr(text_obj, "text", None)
        if text:
            resolved_text = str(text).strip()
            logger.debug(f"_get_text_value_by_ref: resolved '{ref}' from text field -> '{resolved_text}'")
            return resolved_text

        orig = getattr(text_obj, "orig", None)
        if orig:
            resolved_orig = str(orig).strip()
            logger.debug(f"_get_text_value_by_ref: resolved '{ref}' from orig field -> '{resolved_orig}'")
            return resolved_orig

        logger.debug(f"_get_text_value_by_ref: ref '{ref}' resolved to empty text/orig")
    except Exception as e:
        logger.debug(f"_get_text_value_by_ref: failed to resolve ref '{ref}': {e}", exc_info=True)

    return ""


def _looks_like_table_caption(text: str) -> bool:
    """
    Heuristic check for real table captions such as:
    'Table 1-1 VIOS release schedule'
    """
    if not text:
        logger.debug("_looks_like_table_caption: empty text -> False")
        return False
    text_stripped = text.strip()
    is_match = bool(TABLE_CAPTION_PATTERN.match(text_stripped))
    logger.debug(f"_looks_like_table_caption: text='{text_stripped}' match={is_match}")
    return is_match


def _get_ref_value(ref_obj) -> str | None:
    """
    Extract a Docling ref string from dict-like or object-like refs.
    """
    if isinstance(ref_obj, dict):
        return ref_obj.get("$ref")
    return (
        getattr(ref_obj, "ref", None)
        or getattr(ref_obj, "$ref", None)
        or getattr(ref_obj, "cref", None)
        or getattr(ref_obj, "self_ref", None)
    )


def _get_doc_item_by_ref(converted_doc, ref: str):
    """
    Resolve a Docling ref to the underlying object when possible.
    """
    for prefix in ("texts", "tables", "groups", "pictures"):
        idx = _parse_ref_index(ref, prefix)
        if idx is None:
            continue
        try:
            collection = getattr(converted_doc, prefix, None)
            if collection is not None:
                return collection[idx]
        except Exception as e:
            logger.debug(f"_get_doc_item_by_ref: failed to resolve '{ref}' in '{prefix}': {e}", exc_info=True)
            return None
    logger.debug(f"_get_doc_item_by_ref: unsupported or unresolved ref '{ref}'")
    return None


def _get_parent_ref_for_table(converted_doc, table_ix: int) -> str:
    """
    Resolve the parent ref for a table, if any.
    """
    try:
        table_obj = converted_doc.tables[table_ix]
        parent = getattr(table_obj, "parent", None)
        parent_ref = _get_ref_value(parent) if parent is not None else None
        logger.debug(f"_get_parent_ref_for_table: table_ix={table_ix}, parent_ref={parent_ref}")
        return parent_ref or ""
    except Exception as e:
        logger.debug(f"_get_parent_ref_for_table: failed for table_ix={table_ix}: {e}", exc_info=True)
        return ""


def _get_child_refs(item) -> list[str]:
    """
    Return child refs for a Docling item in document order.
    """
    try:
        children = getattr(item, "children", []) or []
        refs = []
        for child in children:
            child_ref = _get_ref_value(child)
            if child_ref:
                refs.append(child_ref)
        return refs
    except Exception as e:
        logger.debug(f"_get_child_refs: failed to extract child refs: {e}", exc_info=True)
        return []


def _find_matching_caption_near_refs(converted_doc, ordered_refs: list[str], target_ref: str, search_window: int) -> str:
    """
    Look for a caption-like text node near the target ref inside an ordered ref list.
    """
    if not ordered_refs:
        logger.debug("_find_matching_caption_near_refs: ordered_refs empty")
        return ""

    try:
        target_pos = ordered_refs.index(target_ref)
        logger.debug(f"_find_matching_caption_near_refs: found {target_ref} at pos={target_pos}")
    except ValueError:
        logger.debug(f"_find_matching_caption_near_refs: target_ref {target_ref} not found in ordered refs")
        return ""

    candidate_positions = list(range(max(0, target_pos - search_window), target_pos))
    candidate_positions.reverse()
    candidate_positions.extend(range(target_pos + 1, min(len(ordered_refs), target_pos + 1 + search_window)))

    candidate_refs = [(pos, ordered_refs[pos]) for pos in candidate_positions]
    logger.debug(f"_find_matching_caption_near_refs: candidate positions/refs={candidate_refs}")

    for pos in candidate_positions:
        ref = ordered_refs[pos]

        if not ref.startswith("#/texts/"):
            logger.debug(f"_find_matching_caption_near_refs: skipping non-text ref {ref}")
            continue

        text = _get_text_value_by_ref(converted_doc, ref)
        logger.debug(f"_find_matching_caption_near_refs: resolved ref {ref} to text='{text}'")

        if _looks_like_table_caption(text):
            logger.debug(f"_find_matching_caption_near_refs: matched caption '{text}' near {target_ref}")
            return text

    logger.debug(f"_find_matching_caption_near_refs: no caption match found near {target_ref}")
    return ""

def _get_enclosing_section_header_for_table(converted_doc, table_ix: int) -> str:
    """
    Secondary fallback for DOCX-like structures where a table is nested under
    a section/container node but has no explicit caption paragraph.
    """
    parent_ref = _get_parent_ref_for_table(converted_doc, table_ix)
    if not parent_ref:
        logger.debug(f"_get_enclosing_section_header_for_table: no parent ref for table_ix={table_ix}")
        return ""

    parent_item = _get_doc_item_by_ref(converted_doc, parent_ref)
    if parent_item is None:
        logger.debug(f"_get_enclosing_section_header_for_table: could not resolve parent item for {parent_ref}")
        return ""

    label = getattr(parent_item, "label", None)
    text = (getattr(parent_item, "text", None) or getattr(parent_item, "orig", None) or "").strip()
    logger.debug(
        f"_get_enclosing_section_header_for_table: table_ix={table_ix}, parent_ref={parent_ref}, "
        f"label={label}, text='{text}'"
    )

    if label == "section_header" and text:
        return text

    return ""


def recover_table_caption_from_body_context(converted_doc, table_ix: int, search_window: int = 3) -> str:
    """
    Recover a table caption using layered fallbacks:
    1. nearby caption paragraph in top-level body order
    2. nearby caption paragraph within the enclosing parent/container children
    3. enclosing section header text as semantic fallback
    """
    target_ref = f"#/tables/{table_ix}"
    logger.debug(f"recover_table_caption_from_body_context: looking for caption near {target_ref} with search_window={search_window}")

    body_refs = _get_body_children_refs(converted_doc)
    caption = _find_matching_caption_near_refs(converted_doc, body_refs, target_ref, search_window)
    if caption:
        logger.debug(f"recover_table_caption_from_body_context: using body-level caption '{caption}' for {target_ref}")
        return caption

    parent_ref = _get_parent_ref_for_table(converted_doc, table_ix)
    if parent_ref:
        parent_item = _get_doc_item_by_ref(converted_doc, parent_ref)
        if parent_item is not None:
            parent_child_refs = _get_child_refs(parent_item)
            caption = _find_matching_caption_near_refs(converted_doc, parent_child_refs, target_ref, search_window)
            if caption:
                logger.debug(
                    f"recover_table_caption_from_body_context: using parent-level nearby caption '{caption}' "
                    f"for {target_ref} within parent {parent_ref}"
                )
                return caption

    section_header = _get_enclosing_section_header_for_table(converted_doc, table_ix)
    if section_header:
        logger.debug(
            f"recover_table_caption_from_body_context: using enclosing section header '{section_header}' "
            f"as secondary fallback for {target_ref}"
        )
        return section_header

    logger.debug(f"recover_table_caption_from_body_context: no caption match found for {target_ref}")
    return ""

def estimate_docx_page_count(docx_path: str) -> int:
    """
    Estimate page count for DOCX based on content and formatting.
    Assumes standard page: ~450 words or ~3000 characters per page.
    
    Args:
        docx_path: Path to the DOCX file
        
    Returns:
        Estimated page count (minimum 1)
    """
    try:
        doc = Document(docx_path)
        
        total_chars = 0
        total_words = 0
        
        # Count paragraph content
        for para in doc.paragraphs:
            text = para.text
            total_chars += len(text)
            total_words += len(text.split())
        
        # Count table content
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text
                    total_chars += len(text)
                    total_words += len(text.split())
        
        # Estimate pages (conservative: 450 words per page)
        estimated_pages = max(1, (total_words // WORDS_PER_PAGE) + 1)
        
        logger.debug(f"DOCX {docx_path}: {total_words} words, estimated {estimated_pages} pages")
        return estimated_pages
        
    except Exception as e:
        logger.error(f"Error estimating page count for {docx_path}: {e}")
        return 1  # Return minimum 1 page on error


def get_docx_toc(docx_path: str) -> Dict[str, int]:
    """
    Extract table of contents from DOCX file.
    
    Strategy (in order of preference):
    1. Try combined extraction (Table Paragraph + List Paragraph) - handles most TOC formats
    2. Fall back to formal TOC styles ('TOC 1', 'TOC 2', etc.)
    3. Fall back to Heading styles ('Heading 1', 'Heading 2', etc.)
    
    Args:
        docx_path: Path to the DOCX file
        
    Returns:
        Dictionary mapping heading text to level (similar to PDF TOC format)
    """
    try:
        # Primary method: Combined extraction (Table Paragraph + List Paragraph)
        toc = extract_toc_combined(docx_path)
        
        if toc:
            logger.info(f"DOCX {docx_path}: extracted {len(toc)} TOC entries from combined extraction")

            logger.debug(f" extracted toc is {toc}")

            return toc
        
        # Fallback 1: Try formal TOC styles
        logger.info(f"DOCX {docx_path}: No combined TOC found, trying TOC styles")
        toc = extract_toc_from_toc_styles(docx_path)
        
        if toc:
            logger.info(f"DOCX {docx_path}: extracted {len(toc)} TOC entries from TOC styles")

            logger.debug(f" extracted toc is {toc}")

            return toc
        
        # Fallback 2: Try Heading styles
        logger.info(f"DOCX {docx_path}: No TOC styles found, trying Heading styles")
        toc = extract_toc_from_headings(docx_path)
        
        if toc:
            logger.info(f"DOCX {docx_path}: extracted {len(toc)} TOC entries from Heading styles")
        else:
            logger.warning(f"DOCX {docx_path}: No TOC found using any method")

        logger.debug(f" extracted toc is {toc}")
        
        return toc
        
    except Exception as e:
        logger.error(f"Error extracting TOC from {docx_path}: {e}")
        return {}

# ============================================================================
# NEW: TOC Style-Based Extraction Functions
# These functions extract TOC from Word's TOC field styles ('TOC 1', 'TOC 2', etc.)
# ============================================================================

def extract_toc_level_from_style(style_name: str) -> int:
    """
    Extract TOC level from Word style name.
    
    Examples:
        'TOC 1' -> 1
        'TOC 2' -> 2
        'toc 3' -> 3
        'TOC Heading' -> 1 (default)
    
    Args:
        style_name: Word paragraph style name
        
    Returns:
        TOC level as integer (defaults to 1 if no number found)
    """
    # Try to extract number from style name
    match = re.search(r'toc\s*(\d+)', style_name, re.IGNORECASE)
    if match:
        return int(match.group(1))
    
    # If no number found (e.g., 'TOC Heading'), default to level 1
    return 1


def extract_toc_from_toc_styles(docx_path: str) -> Dict[str, int]:
    """
    Extract TOC from DOCX file by looking for TOC styles.
    This works ONLY if the document has a formal TOC field with TOC styles.
    
    Word TOC styles are typically named: 'TOC 1', 'TOC 2', 'TOC 3', 'TOC Heading', etc.
    
    Args:
        docx_path: Path to the DOCX file
        
    Returns:
        Dictionary mapping TOC text to level (similar to PDF TOC format)
    """
    try:
        doc = Document(docx_path)
        toc = {}
        
        for paragraph in doc.paragraphs:
            if paragraph.style and paragraph.style.name:
                style_name = paragraph.style.name
                
                # Check for TOC styles: 'TOC 1', 'TOC 2', 'TOC Heading', etc.
                if 'toc' in style_name.lower():
                    text = paragraph.text.strip()
                    if text:
                        # Extract level from style name
                        level = extract_toc_level_from_style(style_name)
                        toc[text] = level
        
        logger.info(f"Extracted {len(toc)} TOC entries from TOC styles in {docx_path}")
        if len(toc) == 0:
            logger.warning(
                f"No TOC styles found in {docx_path}. "
                "Document may not have a formal TOC field. "
                "Consider using extract_toc_from_headings() instead."
            )
        
        return toc
        
    except Exception as e:
        logger.error(f"Error extracting TOC from TOC styles in {docx_path}: {e}")
        return {}


def extract_toc_from_headings(docx_path: str) -> Dict[str, int]:
    """
    Extract TOC from DOCX Heading styles (Heading 1, Heading 2, etc.).
    This is a fallback when document doesn't have formal TOC field.
    
    Args:
        docx_path: Path to the DOCX file
        
    Returns:
        Dictionary mapping heading text to level
    """
    try:
        doc = Document(docx_path)
        toc = {}
        
        for para in doc.paragraphs:
            if para.style and para.style.name.startswith('Heading'):
                text = para.text.strip()
                if text:
                    # Extract level: 'Heading 1' -> 1, 'Heading 2' -> 2
                    match = re.match(r'Heading (\d+)', para.style.name)
                    if match:
                        level = int(match.group(1))
                        toc[text] = level
        
        logger.info(f"Extracted {len(toc)} TOC entries from Heading styles in {docx_path}")
        return toc
        
    except Exception as e:
        logger.error(f"Error extracting TOC from Heading styles in {docx_path}: {e}")
        return {}


def extract_toc_combined(docx_path: str) -> Dict[str, int]:
    """
    Extract TOC from BOTH 'Table Paragraph' and 'List Paragraph' styles.
    This captures TOC entries whether they're in tables or as list items.
    
    This is the recommended approach as it handles various TOC formats:
    - TOC in tables (Table Paragraph style)
    - TOC as lists (List Paragraph style)
    
    Args:
        docx_path: Path to the DOCX file
        
    Returns:
        Dictionary mapping TOC text to level
    """
    try:
        doc = Document(docx_path)
        toc = {}
        
        logger.debug(f"Extracting TOC from Table Paragraphs AND List Paragraphs in {docx_path}")
        
        # Method 1: Extract from tables (Table Paragraph style)
        table_count = 0
        for table_idx, table in enumerate(doc.tables):
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if para.style and para.style.name == 'Table Paragraph':
                            text = para.text.strip()
                            if text:
                                # Skip header row (usually "Contents")
                                if text.lower() in ['contents', 'table of contents']:
                                    continue
                                
                                # Remove page numbers (e.g., "Chapter 1    45" -> "Chapter 1")
                                text_clean = re.sub(r'\s+\d+$', '', text).strip()
                                # Remove trailing dots/leader characters (e.g., "Chapter 1. . . . ." -> "Chapter 1")
                                # Pattern matches dots with spaces between them: ". . . . ." or "....."
                                text_clean = re.sub(r'[\.\s]+$', '', text_clean).strip()
                                
                                if text_clean and text_clean not in toc:
                                    # Skip standalone numbers (likely page numbers or section markers)
                                    if re.match(r'^\d+$', text_clean):
                                        continue
                                    level = _infer_toc_level_from_text(text_clean)
                                    toc[text_clean] = level
                                    table_count += 1
        
        logger.debug(f"Extracted {table_count} entries from Table Paragraph style")

        logger.info(f"Total unique TOC entries extracted: {len(toc)}")

        logger.debug(f"Extracted tocs : {toc}")
        
        return toc
        
    except Exception as e:
        logger.error(f"Error extracting combined TOC from {docx_path}: {e}")
        return {}


def _infer_toc_level_from_text(text: str) -> int:
    """
    Infer TOC level from text content.
    
    Heuristics:
    - "Chapter X" -> level 1
    - "X " (single number) -> level 2
    - "X.Y " (two numbers) -> level 3
    - "X.Y.Z " (three numbers) -> level 4
    - Common sections -> level 1
    - Default -> level 2
    
    Args:
        text: TOC entry text
        
    Returns:
        Inferred level (1-5)
    """
    # Check for chapter
    if text.lower().startswith('chapter '):
        return 1
    
    # Check for numbered sections (e.g., "1.1", "1.2.3")
    match = re.match(r'^(\d+(?:\.\d+)*)\s+', text)
    if match:
        dots = match.group(1).count('.')
        # Correct mapping: "1" (0 dots) -> 2, "1.1" (1 dot) -> 2, "1.2.1" (2 dots) -> 3
        # This ensures "1.1" and "1.2" are at the same level as their parent "1"
        return min(max(dots, 1) + 1, 5)
    
    # Check for common top-level sections
    if text in ['Preface', 'Introduction', 'Contents', 'Notices', 'Trademarks', 'Appendix']:
        return 1
    
    # Default to level 2 for other entries
    return 2
